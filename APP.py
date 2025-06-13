import streamlit as st
import json
import os
from dotenv import load_dotenv # Import load_dotenv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime
import re
import numpy as np
from typing import List, Dict, Any, Tuple
import hashlib
import pickle
from collections import defaultdict
import math
import sqlite3
from pathlib import Path
from io import BytesIO # Import BytesIO for handling file-like objects in memory
from werkzeug.security import generate_password_hash, check_password_hash # For password hashing

# --- Database Setup ---
DB_NAME = "users.db"

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            auth_provider TEXT DEFAULT 'email'
        )
    ''')
    conn.commit()
    conn.close()

# Configure page
st.set_page_config(
    page_title="Algorisk AI",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="collapsed" # Start collapsed for login page
)

# Initialize the database (creates the table if it doesn't exist)
init_db()


# --- Load Environment Variables ---
# This will load variables from a .env file in the same directory or parent directories
# Construct an absolute path to the .env file relative to this script file.
# `__file__` is the path to the current script (App5.py)
script_dir = os.path.dirname(os.path.abspath(__file__))
dotenv_path = os.path.join(script_dir, '.env.txt')
loaded_from_env = load_dotenv(dotenv_path=dotenv_path, override=True) # Override existing system env vars if any conflict, and explicitly pass path



print(f"DEBUG: Attempting to load .env.txt from: {dotenv_path}")
print(f"DEBUG: load_dotenv() returned: {loaded_from_env}")

# --- Gemini API Key Configuration ---
# Get API key from environment variable
import google.generativeai as genai # Import Gemini library
 
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    print(f"DEBUG: Gemini API Key configured (first 5 chars): {GEMINI_API_KEY[:5]}...")
else:
    print("DEBUG: GEMINI_API_KEY not found in environment.")

# Add an early check and display a clear error in Streamlit if the key is missing
if not GEMINI_API_KEY:
    st.error("🔴 CRITICAL ERROR: Gemini API key not found. Please ensure the `GEMINI_API_KEY` is set in your `.env` file (in the same directory as App5.py) or as a system environment variable. The application's AI features will not work without it.")
    # Consider adding st.stop() if the app is entirely unusable without the key,
    # or allow it to run with AI features disabled.

# Custom CSS for the main application (light theme)
MAIN_APP_CSS = """
<style>
    /* Styles for the main application after login */
    /* Core Dark Theme */
    body { /* Light theme body */
        color: #333333; /* Dark text for light background */
        background-color: #f0f2f6; /* Light grey background for the entire page */
    }

    .main {
        background-color: #f0f2f6; /* Light grey background for main content area */
        color: #333333; /* Dark text */
    }
    /* Streamlit's main content area */
    .block-container {
        background-color: #FFFFFF; /* White card background */
        border: 1px solid #d1d1d1; /* Light grey border */
        border-radius: 15px;
        padding: 2rem;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05); /* Subtle shadow */
        margin: 1rem 0;
        color: #333333; /* Dark text within cards */
    }
   /* Main App Headings - should be dark for light background */
    h1, h2, h3, h4, h5, h6 {
       
        color: #333333; /* Dark color for headings in the main app */
    }
    .main .block-container h1 { /* More specific for main app title if needed */
         color: #333333 !important;
    }
    h1 {
        text-align: center;
        font-size: 2.8em !important;
        font-weight: bold;
        margin-bottom: 0.2em !important;
    }
    /* Subtext for hero */
    .hero-subtext {
        text-align: center;
        color: #555555; /* Darker grey for subtext */
        font-size: 1.2em;
        margin-bottom: 2em;
    }

    /* Buttons */
    .stButton > button {
        background-color: #007bff; /* Standard blue for buttons */
        color: #FFFFFF; /* White text on buttons */
        border: 1px solid #007bff;
        border-radius: 8px;
        padding: 0.6rem 1.2rem;
        font-weight: bold;
        transition: all 0.3s ease;
        /* box-shadow: 0 0 5px rgba(0, 123, 255, 0); /* Optional: remove or adjust shadow */
    }
    .stButton > button:hover {
        background-color: #0056b3; /* Darker blue on hover */
        border-color: #0056b3;
        /* box-shadow: 0 0 10px rgba(0, 123, 255, 0.5); /* Optional: adjust hover shadow */
        transform: translateY(-2px);
    }
    /* Primary action button */
    .stButton.primary-action-button > button {
        background: linear-gradient(45deg, #007bff, #0056b3); /* Blue gradient */
        color: #FFFFFF;
        font-size: 1.1em;
        padding: 0.8rem 1.5rem;
    }
    .stButton.primary-action-button > button:hover {
        box-shadow: 0 0 15px rgba(0, 123, 255, 0.6); /* Adjusted shadow for light theme */
    }

    /* Input fields and text areas */
    .stTextInput > div > div > input, .stTextArea > div > div > textarea {
        background-color: #FFFFFF; /* White background for inputs */
        color: #333333; /* Dark text in inputs */
        border: 1px solid #ced4da; /* Standard light border */
        border-radius: 5px;
    }
    .stFileUploader > div > div > button { /* Style file uploader button */
        border: 1px dashed #007bff; /* Blue dashed border */
        background-color: rgba(0, 123, 255, 0.05); /* Very light blue background */
        color: #007bff; /* Blue text */
    }
    .stFileUploader > div > div:hover { /* Glowing outline for dropzone */
        border-color: #007bff;
        box-shadow: 0 0 8px rgba(0, 123, 255, 0.3); /* Softer glow */
    }

    /* Selectbox and Multiselect */
    .stSelectbox > div > div, .stMultiSelect > div > div {
        background-color: #FFFFFF; /* White background */
        color: #333333; /* Dark text */
        border: 1px solid #ced4da; /* Light border, if not handled by Streamlit default */
    }

    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        background-color: #f0f2f6; /* Light background for tab list */
        border-bottom: 2px solid #007bff; /* Blue bottom border for active tab underline */
    }
    .stTabs [data-baseweb="tab"] {
        background-color: #f0f2f6; /* Light background for inactive tabs */
        color: #555555; /* Dark grey text for inactive tabs */
        font-weight: bold;
    }
    .stTabs [data-baseweb="tab--selected"] {
        background-color: #FFFFFF; /* White background for selected tab */
        color: #007bff; /* Blue text for selected tab */
        border-bottom: 2px solid #007bff; /* Consistent blue underline */
    }
/* Sidebar */
.css-1d391kg { /* Streamlit sidebar class */
    background-color: #f9f9f9; /* Light grey for a clean sidebar */
    border-right: 1px solid #dcdcdc; /* Subtle light border */
}
.css-1d391kg {
    background-color: #ffffff;  /* Pure white background */
    border-right: 1px solid #e0e0e0; /* Soft, modern border */
    box-shadow: 2px 0 4px rgba(0, 0, 0, 0.05); /* Light shadow for depth */
}

    .css-1d391kg .stMarkdown, .css-1d391kg .stSelectbox, .css-1d391kg .stSlider {
        color: #333333; /* Dark text in sidebar */
    }

    /* Horizontal Ticker */
    .ticker-wrap {
        width: 100%;
        overflow: hidden;
        background: rgba(220, 220, 230, 0.9); /* Light translucent background */
        border: 1px solid #007bff; /* Blue border */
        border-radius: 8px;
        padding: 0.5em 0;
        margin-bottom: 2em;
        box-shadow: 0 0 10px rgba(0, 123, 255, 0.2); /* Softer blue shadow */
    }
    .ticker {
        display: inline-block;
        white-space: nowrap;
        padding-left: 100%;
        animation: ticker-scroll 30s linear infinite;
        color: #333333; /* Dark text for ticker */
    }
    .ticker span {
        display: inline-block;
        padding: 0 2rem; /* Spacing between items */
        font-size: 0.9em;
    }
    @keyframes ticker-scroll {
        0% { transform: translateX(0); }
        100% { transform: translateX(-100%); }
    }

    /* Footer */
    .footer {
        text-align: center;
        padding: 1.5em 0;
        color: #555555; /* Dark grey for footer text */
        font-size: 0.9em;
        border-top: 1px solid #e0e0e0; /* Light grey top border */
        margin-top: 2em;
    }

    /* Styling for RAG workflow steps (example) */
    .rag-step, .rag-success, .rag-processing, .rag-error {
        background: #FFFFFF; /* White background for steps */
        padding: 1rem;
        margin: 0.5rem 0;
        border-radius: 0 8px 8px 0;
        color: #333333; /* Dark text for steps */
    }
    .rag-step { border-left: 4px solid #007bff; } /* Blue for info */
    .rag-success { border-left: 4px solid #28a745; } /* Green for success */
    .rag-processing { border-left: 4px solid #ffc107; } /* Yellow for processing */
    .rag-error { border-left: 4px solid #dc3545; } /* Red for error */

    .workflow-container { /* For RAG workflow display */
        background: #FFFFFF; /* White background */
        border: 1px solid #d1d1d1; /* Light grey border */
        border-radius: 10px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    .chunk-preview {
        background: #f8f9fa; /* Light grey, good for code/text previews */
        border: 1px solid #dee2e6; /* Standard light border */
        border-radius: 8px;
        padding: 0.5rem;
        margin: 0.25rem 0;
        font-family: monospace;
        font-size: 0.8em;
        color: #212529; /* Darker text for previews */
    }
    /* Ensure auth form elements on the dark landing page are styled appropriately */
    .auth-container-dark {
        background-color: rgba(17, 24, 39, 0.8); /* Dark, slightly transparent background for form card */
        padding: 2rem;
        border-radius: 12px;
        border: 1px solid rgba(249, 115, 22, 0.2); /* Subtle orange border */
        box-shadow: 0 8px 25px rgba(0, 0, 0, 0.3);
    }
    .auth-container-dark h1, .auth-container-dark h2, .auth-container-dark h3 {
        color: #FFFFFF !important; /* White headings for dark form */
    }
    .auth-container-dark .stTextInput > div > div > input,
    .auth-container-dark .stTextArea > div > div > textarea {
        background-color: #1f2937; /* Dark input background */
        color: #e5e7eb; /* Light text in inputs */
        border: 1px solid #4b5563; /* Grey border */
    }
    .auth-container-dark .stButton > button {
        background-color: #f97316; /* Orange for buttons */
        color: #FFFFFF;
        border: 1px solid #f97316;
    }
    .auth-container-dark .stButton > button:hover {
        background-color: #ea580c; /* Darker orange on hover */
        border-color: #ea580c;
    }
    .auth-container-dark .stButton > button:not([kind="primary"]) { /* For secondary buttons like "Sign Up" / "Login" toggles */
        background: transparent;
        color: #fdba74; /* Lighter orange for text */
        border: 1px solid #f97316;
    }
     .auth-container-dark .stButton > button:not([kind="primary"]):hover {
        background-color: rgba(249, 115, 22, 0.1);
        color: #f97316;
    }
    </style>
"""
st.markdown(MAIN_APP_CSS, unsafe_allow_html=True)

# =============================================
# RAG FRAMEWORK IMPLEMENTATION
# =============================================

class SimpleEmbedding:
    """Simple TF-IDF based embedding for RAG without external dependencies"""
    
    def __init__(self):
        self.vocabulary = {}
        self.idf = {}
        self.documents = []
        
    def _tokenize(self, text: str) -> List[str]:
        """Simple tokenization"""
        text = re.sub(r'[^\w\s]', ' ', text.lower())
        return [word for word in text.split() if len(word) > 2]
    
    def _compute_tf(self, tokens: List[str]) -> Dict[str, float]:
        """Compute term frequency"""
        tf = defaultdict(int)
        for token in tokens:
            tf[token] += 1
        
        # Normalize by document length
        max_freq = max(tf.values()) if tf else 1
        return {term: freq/max_freq for term, freq in tf.items()}
    
    def fit(self, documents: List[str]):
        """Fit the embedding model on documents"""
        self.documents = documents
        
        # Build vocabulary and document frequency
        doc_freq = defaultdict(int)
        all_tokens = []
        
        for doc in documents:
            tokens = self._tokenize(doc)
            all_tokens.append(tokens)
            unique_tokens = set(tokens)
            for token in unique_tokens:
                doc_freq[token] += 1
        
        # Build vocabulary
        self.vocabulary = {token: idx for idx, token in enumerate(doc_freq.keys())}
        
        # Compute IDF
        num_docs = len(documents)
        self.idf = {
            token: math.log(num_docs / freq) 
            for token, freq in doc_freq.items()
        }
        
        return all_tokens
    
    def transform(self, documents: List[str]) -> np.ndarray:
        """Transform documents to TF-IDF vectors"""
        vectors = []
        
        for doc in documents:
            tokens = self._tokenize(doc)
            tf = self._compute_tf(tokens)
            
            # Create TF-IDF vector
            vector = np.zeros(len(self.vocabulary))
            for token, tf_val in tf.items():
                if token in self.vocabulary:
                    idx = self.vocabulary[token]
                    vector[idx] = tf_val * self.idf.get(token, 0)
            
            vectors.append(vector)
        
        return np.array(vectors)
    
    def cosine_similarity(self, vec1: np.ndarray, vec2: np.ndarray) -> float:
        """Compute cosine similarity between two vectors"""
        dot_product = np.dot(vec1, vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        
        if norm1 == 0 or norm2 == 0:
            return 0
        
        return dot_product / (norm1 * norm2)

class DocumentChunker:
    """Document chunking with overlap for better retrieval"""
    
    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def chunk_text(self, text: str, metadata: Dict = None) -> List[Dict]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            chunk_data = {
                'text': chunk_text,
                'start_idx': i,
                'end_idx': min(i + self.chunk_size, len(words)),
                'metadata': metadata or {}
            }
            chunks.append(chunk_data)
            
            if i + self.chunk_size >= len(words):
                break
        
        return chunks

class VectorStore:
    """Simple vector store for document retrieval"""
    
    def __init__(self, embedding_model: SimpleEmbedding):
        self.embedding_model = embedding_model
        self.documents = [] # Stores chunk texts
        self.vectors = None # Stores chunk vectors
        self.metadata = [] # Stores chunk metadata
    
    def add_documents(self, texts: List[str], metadata: List[Dict] = None):
        """Add documents (chunks) to the vector store"""
        self.documents.extend(texts)
        self.metadata.extend(metadata or [{}] * len(texts))
        
        # Fit embedding model and transform documents
        # Re-fit the model on the cumulative documents each time
        self.embedding_model.fit(self.documents)
        self.vectors = self.embedding_model.transform(self.documents)
    
    def similarity_search(self, query: str, k: int = 5) -> List[Dict]:
        """Search for similar documents"""
        if self.vectors is None or len(self.vectors) == 0:
            return []
        
        # Transform query
        query_vector = self.embedding_model.transform([query])[0]
        
        # Compute similarities
        similarities = []
        for i, doc_vector in enumerate(self.vectors):
            sim = self.embedding_model.cosine_similarity(query_vector, doc_vector)
            similarities.append((sim, i))
        
        # Sort by similarity and return top k
        similarities.sort(reverse=True, key=lambda x: x[0])
        
        results = []
        for sim_score, idx in similarities[:k]:
            results.append({
                'text': self.documents[idx],
                'metadata': self.metadata[idx],
                'similarity': sim_score
            })
        
        return results

class SimpleRAG:
    """Complete RAG implementation with workflow tracking"""
    
    def __init__(self):
        self.embedding_model = SimpleEmbedding()
        # Initialize chunker and vector store here, but update settings later
        self.chunker = DocumentChunker() 
        self.vector_store = VectorStore(self.embedding_model)
        self.knowledge_base = [] # Stores original documents
        self.workflow_steps = []
    
    def log_step(self, step: str, status: str = "info", details: str = ""):
        """Log workflow step"""
        self.workflow_steps.append({
            'step': step,
            'status': status,
            'details': details,
            'timestamp': datetime.now().strftime("%H:%M:%S")
        })
    
    def build_knowledge_base(self, documents: List[Dict], progress_callback=None):
        """Build knowledge base from documents"""
        self.log_step("🚀 Starting Knowledge Base Construction", "info")
        self.knowledge_base = documents # Store original documents
        
        if progress_callback:
            progress_callback(10, "Processing documents...")
        
        all_chunks = []
        all_metadata = []
        
        for i, doc in enumerate(documents):
            self.log_step(f"📄 Processing document: {doc.get('source', f'Document {i+1}')}", "processing")
            
            # Chunk document using the current chunker settings
            chunks = self.chunker.chunk_text(doc['content'], {
                'source': doc.get('source', f'doc_{i}'),
                'type': doc.get('type', 'unknown')
            })
            
            chunk_texts = [chunk['text'] for chunk in chunks]
            chunk_metadata = [chunk['metadata'] for chunk in chunks]
            
            all_chunks.extend(chunk_texts)
            all_metadata.extend(chunk_metadata)
            
            self.log_step(f"✅ Created {len(chunks)} chunks from {doc.get('source', f'Document {i+1}')}", "success")
            
            if progress_callback:
                progress_callback(10 + (i + 1) * 60 // len(documents), f"Processed {i+1}/{len(documents)} documents")
        
        self.log_step(f"🔧 Building vector embeddings for {len(all_chunks)} chunks", "processing")
        
        if progress_callback:
            progress_callback(75, "Building vector embeddings...")
        
        # Add to vector store
        # Re-initialize vector store to clear previous data if rebuilding
        self.vector_store = VectorStore(self.embedding_model)
        self.vector_store.add_documents(all_chunks, all_metadata)
        
        self.log_step(f"✅ Knowledge base built successfully with {len(all_chunks)} chunks", "success")
        
        if progress_callback:
            progress_callback(100, "Knowledge base construction complete!")
        
        return len(all_chunks)
    
    def retrieve_context(self, query: str, k: int = 5) -> List[Dict]:
        """Retrieve relevant context for a query"""
        self.log_step(f"🔍 Retrieving context for query: '{query[:50]}...'", "processing")
        
        results = self.vector_store.similarity_search(query, k=k)
        
        self.log_step(f"✅ Retrieved {len(results)} relevant chunks", "success", 
                     f"Top similarity: {results[0]['similarity']:.3f}" if results else "No results found")
        
        return results
    
    def generate_response(self, query: str, context: List[Dict], ai_model: str = "gpt-4", temperature: float = 0.3, max_tokens: int = 1500) -> str:
        """Generate response using retrieved context (simplified version)"""
        self.log_step("🤖 Generating AI response with retrieved context", "processing")
        
        if not context and not query:
            return "I don't have sufficient information in the knowledge base or a query to answer."
        
        # Using LLM for a more sophisticated response
        if context:
            context_text = "\n---\n".join([item['text'] for item in context])
            prompt_message = f"""
You are a helpful AI assistant with access to a knowledge base.
Answer the following question based ONLY on the provided context.
If you cannot answer the question based on the context, say "I cannot answer this question based on the provided information."

Context:
{context_text}

Question: {query}

Answer:
"""
            try:
                if not GEMINI_API_KEY:
                    self.log_step("Gemini API key not configured.", "error")
                    return "Error: Gemini API key not configured. Please set the GEMINI_API_KEY environment variable."

                # Map ai_model to a valid Gemini model name if necessary
                gemini_model_name = ai_model
                # Let's default to a known good model if the selected one is problematic or an old OpenAI name
                if "gpt" in gemini_model_name.lower() or gemini_model_name == "gemini-pro": 
                    gemini_model_name = "gemini-1.5-flash-latest" # Or try "gemini-1.0-pro"
                
                model_instance = genai.GenerativeModel(gemini_model_name)
                
                generation_config = genai.types.GenerationConfig(temperature=temperature, max_output_tokens=max_tokens)
                llm_response = model_instance.generate_content(prompt_message, generation_config=generation_config)
                response = llm_response.text
            except Exception as e:
                self.log_step(f"Error generating LLM response: {e}", "error")
                response = f"Error generating LLM response: {e}"
        else: # Fallback if no context but query exists (though RAG usually provides context for relevant queries)
            response = "No specific context was retrieved from the knowledge base to answer your query. Please try rephrasing or ensure the knowledge base contains relevant information."

        # Ensure the response is a string
        if not isinstance(response, str):
            self.log_step("Generated response is not a string, converting.", "info")
            response = str(response)

        self.log_step("✅ Response generated successfully", "success")
        
        return response
    
    def query(self, question: str, k: int = 5) -> Dict:
        """Complete RAG query with workflow tracking"""
        self.workflow_steps = []  # Reset workflow for a new query
        
        self.log_step("🎯 Starting RAG Query Process", "info", f"Query: {question}")
        
        # Retrieve context
        context = self.retrieve_context(question, k=k)
        
        # Generate response
        # Use AI settings from session state
        response = self.generate_response(
            question, 
            context,
            st.session_state.get('ai_model', 'gpt-4'), 
            st.session_state.get('temperature', 0.3), 
            st.session_state.get('max_tokens', 1500)
        )
        
        self.log_step("🎉 RAG Query Process Complete", "success")
        
        return {
            'question': question,
            'response': response,
            'context': context,
            'workflow': self.workflow_steps
        }

# =============================================
# VALIDATION CHALLENGES & SCORING
# =============================================

# Validation Challenges List (Comprehensive)
validation_challenges = {
    "A. Regulatory & Governance Challenges": [
        "Misalignment with PRA SS1/23 expectations",
        "Failure to demonstrate forward-looking information integration",
        "Incomplete documentation on model assumptions",
        "Lack of evidence of internal challenge/independent validation",
        "Weak justification for staging thresholds (SICR triggers)",
        "Inadequate validation of staging logic and override policy",
        "Insufficient backtesting of modelled PDs vs actual defaults",
        "Absence of model governance traceability and sign-offs",
        "No model risk tiering and prioritization defined",
        "Failure to separate IFRS 9 and IRB objectives clearly"
    ],
    "B. Data & Input Issues": [
        "Missing or incomplete historical data for key drivers",
        "Data not aligned with reporting/accounting granularity",
        "Biased datasets (e.g. selection bias due to charge-offs)",
        "Ignoring payment holiday or forbearance periods",
        "Insufficient cure window post-default",
        "Inadequate performance window definitions",
        "No vintage or cohort tracking in raw data",
        "Errors in macroeconomic input mapping (lag, level)",
        "Inconsistent treatment of revolving vs transactors",
        "Incorrect use of behavioural data (e.g. utilization)"
    ],
    "C. Model Methodology Challenges": [
        "Logistic regression lacks explanatory power (low Gini)",
        "Model performance degrades materially across vintages",
        "Lifetime PD extension lacks statistical justification",
        "Use of survival model lacks cure-adjusted hazard rates",
        "No explicit error bounds/confidence intervals provided",
        "Ignoring maturity effect on credit card behaviour",
        "Transition matrix method misapplied or unstable",
        "Spurious correlation between macro drivers and defaults",
        "Overfitting in machine learning models (XGBoost, RF)",
        "Ignoring account seasoning in PD maturity profile"
    ],
    "D. Assumptions & Simplifications": [
        "Lifetime PD extrapolation is linear/unrealistic",
        "Too few macroeconomic scenarios used",
        "Macro scenarios not severe enough for downturn risk",
        "No sensitivity analysis for assumptions",
        "Ignoring behavioural changes due to pricing/APR changes",
        "Static exposure profiles assumed over life",
        "No borrower-level dynamic modelling",
        "Lack of justification for segmentation rules",
        "Default definition differs from policy/international norms",
        "Proxy used for default where DPD not reliable"
    ],
    "E. Macroeconomic Integration Weaknesses": [
        "Weak or no statistical linkage between macro and PD",
        "Scenario expansion method (e.g., delta shift) is overly simplistic",
        "Ignoring interaction effects among macro variables",
        "No validation of macroeconomic overlay adjustments",
        "Manual scenario overrides not tracked or justified",
        "Misaligned macroeconomic scenario time horizon",
        "Failure to document macroeconomic source credibility",
        "Over-dependence on GDP/unemployment without justification",
        "No base-vs-downturn impact comparison",
        "Model fails to show responsiveness to stressed scenarios"
    ],
    "F. Portfolio-Specific Limitations": [
        "Model does not differentiate between transactors and revolvers",
        "Lack of segmentation for teaser-rate or 0% BT products",
        "Fails to model exposure dynamics for dormant/reactivated accounts",
        "Lifetime PD doesn’t reflect prepayment/closure risks",
        "Changes in credit limits and exposure not accounted for",
        "No vintage sensitivity tracking over time",
        "Mismatch between booked exposures and active accounts",
        "Product-specific risk features not captured (e.g. loyalty cards)",
        "Ignoring self-cures and behavioural cyclicality",
        "Model not recalibrated post COVID/shock period"
    ],
    "G. Model Validation Weaknesses": [
        "No independent challenger model developed",
        "Validation sample too short or not representative",
        "Inadequate out-of-time testing",
        "Stability metrics (PSI, CoV) not monitored or reported",
        "KS/AUC not benchmarked to acceptable thresholds",
        "Lack of lift analysis across PD bands",
        "Default prediction errors not diagnosed across segments",
        "PD distributions do not match expected shape",
        "Rejection inference not addressed (origination model reused)",
        "Benchmarking with peer models or vendor models missing"
    ],
    "H. Calibration & Monitoring Issues": [
        "No lifetime calibration back to observed defaults",
        "Transition from 12m to lifetime PD poorly controlled",
        "Calibration overrides undocumented or unjustified",
        "No cap/floor logic on PDs for rare segments",
        "Failure to update model in light of monitoring breaches",
        "Model recalibration frequency unclear or ad hoc",
        "Output drifts without input change not explained",
        "Lack of automated monitoring or alerting system",
        "Model decay not assessed annually",
        "Missing override logs and escalation steps"
    ]
}

# =============================================
# HELPER FUNCTIONS
# =============================================

def extract_text_from_file(file):
    """Extract text from uploaded file"""
    # python-docx is needed for .docx, ensure it's installed: pip install python-docx
    try:
        from docx import Document as DocxDocument # Alias to avoid conflict
    except ImportError:
        DocxDocument = None
        st.warning("`python-docx` library not found. DOCX parsing will be limited. Install with `pip install python-docx`.")

    try:
        file.seek(0) # Reset file pointer
        if file.type == "text/plain":
            return file.read().decode('utf-8', errors='ignore')
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            if DocxDocument:
                try:
                    doc = DocxDocument(BytesIO(file.read()))
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return text if text else f"Could not extract text from DOCX: {file.name}"
                except Exception as e:
                    return f"Error processing DOCX {file.name}: {str(e)}"
            else:
                return f"DOCX file ({file.name}) uploaded, but `python-docx` is not available for parsing."
        elif file.type == "application/pdf":
             # Placeholder for PDF processing (e.g., using PyPDF2 or pdfminer.six)
             # For now, return a message indicating PDF processing is not fully implemented.
             # To implement: pip install PyPDF2
             # from PyPDF2 import PdfReader
             # reader = PdfReader(BytesIO(file.read()))
             # text = ""
             # for page in reader.pages:
             #     text += page.extract_text() + "\n"
             # return text
             return f"PDF content from {file.name} (PDF processing not fully implemented in this example)"
        else:
             # Attempt to read as text for other types
             file.seek(0)
             return file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        return f"Error reading file {file.name}: {str(e)}"

def get_builtin_ifrs9_knowledge():
    """Provides hardcoded or loaded built-in IFRS 9 knowledge."""
    # In a real application, this would load from files or a database.
    # For this example, we'll use placeholder text.
    return [
        {'content': "IFRS 9 requires entities to recognise a loss allowance for expected credit losses (ECL) on financial assets. This involves a three-stage approach based on changes in credit quality since initial recognition.", 'source': 'IFRS 9 Summary', 'type': 'regulatory'},
        {'content': "A significant increase in credit risk (SICR) triggers the movement of a financial instrument from Stage 1 (12-month ECL) to Stage 2 (lifetime ECL). SICR assessment should be forward-looking and consider all reasonable and supportable information.", 'source': 'IFRS 9 Staging', 'type': 'regulatory'},
        {'content': "Lifetime PD is the probability of default occurring over the expected life of the financial instrument. It should reflect current conditions and reasonable and supportable forecasts of future economic conditions.", 'source': 'IFRS 9 PD', 'type': 'regulatory'},
        {'content': "Model documentation should be comprehensive, covering purpose, scope, data, methodology, validation, limitations, governance, and monitoring. It must be sufficient for a knowledgeable third party to understand the model's design and operation.", 'source': 'Regulatory Doc Template', 'type': 'model_doc'},
        {'content': "Validation includes assessing model performance (e.g., AUC, Gini, KS), stability (e.g., PSI), and compliance with regulatory requirements. Backtesting and benchmarking are crucial components.", 'source': 'Model Validation Guide', 'type': 'regulatory'},
        {'content': "PRA SS1/23 sets out the Prudential Regulation Authority's expectations for model risk management for firms. It covers model development, implementation, use, validation, and governance.", 'source': 'PRA SS1/23 Overview', 'type': 'regulatory'},
    ]

def analyze_code_structure(code_content):
    """Analyze the structure and components of the uploaded code."""
    analysis = {
        'summary': f"Code consists of {len(code_content.splitlines())} lines.", # Basic summary
        'complexity': {
            'lines': len(code_content.split('\n')),
            'functions': len(re.findall(r'def\s+\w+\s*\(', code_content)), # More robust function regex
            'classes': len(re.findall(r'class\s+\w+\s*[:\(]', code_content)), # More robust class regex
            'imports': len(re.findall(r'^\s*import\s+|^from\s+', code_content, re.MULTILINE))
        },
        'features': [], # Placeholder, can be enhanced with AST parsing
        'libraries': list(set([match[0] or match[1] for match in re.findall(r'import\s+([\w\.]+)|from\s+([\w\.]+)', code_content)])), # Unique libraries
        'functions_list': re.findall(r'def\s+(\w+)\s*\(', code_content) # List of function names
    }
    
    # Extract potential model features (simple regex, can be improved)
    feature_patterns = [
        r'\b(age|income|balance|utilization|score|risk|ltv|dti)\b', # Common financial terms
        r'X_train\[\'"[\'"]\]', # Features from X_train assignments
        r'data\[\'"[\'"]\]' # Features from dataframe access
    ]
    
    identified_features = set()
    for pattern in feature_patterns:
        matches = re.findall(pattern, code_content, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple): # Handle multiple capture groups
                for m_part in match:
                    if m_part and len(m_part) > 2: identified_features.add(m_part.lower())
            elif match and len(match) > 2:
                identified_features.add(match.lower())
    analysis['features'] = sorted(list(identified_features))
    
    return analysis

# Define the list of sections globally
sections = [
    "Executive Summary", "Model Purpose and Use", "Portfolio Segmentation",
    "Data Inputs and Sources", "Missing Data and Imputation Handling", "Model Methodology",
    "Feature Engineering and Variable Selection", "Macroeconomic Linkage (PIT models)",
    "Model Training and Validation Strategy", "Model Performance Results",
    "Backtesting and Benchmarking", "Overlays and Expert Judgment",
    "Model Governance and Lifecycle", "Monitoring Framework",
    "Assumptions and Limitations", "Regulatory Compliance Assessment",
    "Deployment and Implementation", "Audit Trail and Documentation Traceability"
]

def generate_documentation_content(code_analysis, results, model_type, portfolio_type, regulations, code_content_str=None, data_content_str=None, config_content_str=None):
    """Generate comprehensive documentation content (Standard version)."""
    
    # Use provided content or placeholders
    code_summary = code_analysis.get('summary', 'Analysis not available.')
    data_summary = data_content_str[:500] + "..." if data_content_str and len(data_content_str) > 500 else (data_content_str if data_content_str else "No data file uploaded.")
    config_summary = config_content_str[:500] + "..." if config_content_str and len(config_content_str) > 500 else (config_content_str if config_content_str else "No config file uploaded.")

    # This is a simplified version. A real LLM call would be more dynamic.
    # For demonstration, we'll use f-strings with placeholders.
    # In a real scenario, each section's content would be generated by an LLM call
    # similar to generate_rag_enhanced_documentation but without the RAG context.

    doc_content = {}
    for section_title_item in sections: # Use the global 'sections' list
        doc_content[section_title_item] = f"""
        This is the placeholder content for the '{section_title_item}' section.
        Model Type: {model_type}
        Portfolio Type: {portfolio_type}
        Applicable Regulations: {', '.join(regulations)}

        Code Analysis Summary: {code_analysis.get('summary', 'N/A')}
        Identified Libraries: {', '.join(code_analysis.get('libraries', ['N/A'])[:5])}
        Model Performance (AUC): {results.get('auc', 'N/A')}

        Further details for this section would be generated by an AI model based on the inputs.
        Data Summary: {data_summary}
        Config Summary: {config_summary}
        """
    return doc_content

def generate_rag_enhanced_documentation(rag_system: SimpleRAG, code_analysis: Dict, results: Dict, model_type: str, portfolio_type: str, regulations: List[str], code_content: str, data_content: str, config_content: str, additional_info: str, workflow_placeholder):
    """
    Generates documentation content for each section using RAG and LLM.
    """
    doc_content = {}
    # Use the global 'sections' list

    # Summarize data and config for context
    data_summary = data_content[:500] + "..." if data_content and len(data_content) > 500 else (data_content if data_content else "No data file uploaded.")
    config_summary = config_content[:500] + "..." if config_content and len(config_content) > 500 else (config_content if config_content else "No config file uploaded.")

    for i, section_item in enumerate(sections): # Use 'sections' list
        with workflow_placeholder.container():
            st.markdown(f'<div class="rag-processing">🤖 Generating section: **{section_item}**...</div>', unsafe_allow_html=True)

        # Use RAG to retrieve context relevant to the section and model specifics
        rag_query = f"IFRS 9 {model_type} model documentation section: {section_item}. Portfolio: {portfolio_type}. Regulations: {', '.join(regulations)}"
        retrieved_context = rag_system.retrieve_context(rag_query, k=st.session_state.get('retrieval_k', 5))
        context_text = "\n---\n".join([item['text'] for item in retrieved_context])

        prompt = f"""
You are an expert regulatory model developer writing bank and audit-level documentation for an IFRS 9 model.
Your goal is to create precise, auditable, and defensible content for the specified section, leveraging the provided context.
Do NOT hallucinate information not supported by the context or uploaded files.

Write the following section:
- Section: {section_item}
- Model Type: {model_type}
- Portfolio Type: {portfolio_type}
- Regulations: {', '.join(regulations)}

Context from Uploaded Files:
Model Code Summary: {code_analysis.get('summary', 'Analysis not available.')}
Model Code Snippets: {code_content[:1000]}... # Include start of code
Model Results: {json.dumps(results, indent=2)}
Sample Data Snippet/Summary: {data_summary}
Configuration Snippet/Summary: {config_summary}
Additional Info: {additional_info}

Context from Knowledge Base (RAG):
{context_text}

Instructions:
1. Write the content for the "{section_item}" section.
2. Ensure the content is relevant to the Model Type, Portfolio Type, and Regulations.
3. Directly reference or interpret information from the "Context from Uploaded Files" where applicable (e.g., mention specific libraries from code analysis, reference results).
4. Incorporate insights and requirements from the "Context from Knowledge Base (RAG)".
5. Maintain a formal, regulatory, and auditable tone.
6. If information required for a subsection is missing from the context, state that the information was not provided and needs to be supplied.
7. Do NOT make up technical details about the model's implementation if they are not in the code or context.
"""
        try:
            if not GEMINI_API_KEY:
                 raise ValueError("Gemini API key not configured.")

            gemini_model_name = st.session_state.get('ai_model', 'gemini-pro')
            # Default to a known working model if the selection is an old OpenAI model or potentially problematic
            if "gpt" in gemini_model_name.lower() or gemini_model_name == "gemini-pro":
                gemini_model_name = "gemini-1.5-flash-latest" # Or try "gemini-1.0-pro"
            model_instance = genai.GenerativeModel(gemini_model_name)
            
            generation_config = genai.types.GenerationConfig(
                max_output_tokens=st.session_state.get('max_tokens', 1500),
                temperature=st.session_state.get('temperature', 0.3)
            )
            llm_response = model_instance.generate_content(prompt, generation_config=generation_config)
            doc_content[section_item] = llm_response.text
            with workflow_placeholder.container():
                 st.markdown(f'<div class="rag-success">✅ Section **{section_item}** generated.</div>', unsafe_allow_html=True)
        except Exception as e:
            doc_content[section_item] = f"Error generating section '{section_item}': {e}. Please ensure your Gemini API key is correctly configured and the model name is valid."
            with workflow_placeholder.container():
                 st.markdown(f'<div class="rag-error">❌ Error generating section **{section_item}**: {e}</div>', unsafe_allow_html=True)

        # Assuming workflow_placeholder is a Streamlit element that supports progress updates
        if hasattr(workflow_placeholder, 'progress'):
            workflow_placeholder.progress((i + 1) / len(sections))

    return doc_content

def create_word_document(doc_content: Dict[str, str], model_type: str) -> Tuple[BytesIO, str]:
    """Create a professional Word document with the generated content."""
    # Imports moved to top level

    doc = Document()
    
    # Title page
    title = doc.add_heading(f'IFRS 9 {model_type} Model Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('AlgoRisk AI Platform')
    doc.add_page_break()
    
    # Table of contents placeholder (can be updated later or left as placeholder)
    doc.add_heading('Table of Contents', level=1)
    # Simple placeholder list
    for i, section_title_item in enumerate(doc_content.keys(), 1):
         doc.add_paragraph(f'{i}. {section_title_item}')
    doc.add_page_break()
    
    # Add sections
    for section_title_item, content_item in doc_content.items():
        doc.add_heading(section_title_item, level=1)
        doc.add_paragraph(content_item.strip())
        doc.add_paragraph()  # Add spacing
    
    # Save document to BytesIO
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    # Return the BytesIO object and a suggested filename
    file_name = f"ifrs9_{model_type.lower().replace(' (probability of default)', '_pd').replace(' (loss given default)', '_lgd').replace(' (exposure at default)', '_ead')}_documentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    return doc_buffer, file_name

def run_rag_validation_assessment(rag_system: SimpleRAG | None, code_analysis: Dict, validation_level: str, doc_text: str, code_content: str):
    """
    Runs a RAG-enhanced validation assessment against challenges.
    Includes scoring and risk rating.
    """
    assessment_results = {
        "overall_score": 0,
        "total_possible_score": 0,
        "compliant_count": 0,
        "partially_compliant_count": 0,
        "non_compliant_count": 0,
        "error_count": 0,
        "sections": {},
        "gap_table": [],
        "risk_rating": "N/A",
        "summary": ""
    }

    # Scoring weights (can be adjusted)
    rating_scores = {"Compliant": 3, "Partially Compliant": 1, "Non-Compliant": 0, "Error": 0}
    # Use the global validation_challenges dictionary
    section_weights = {section_item: 1 for section_item in validation_challenges.keys()} # Equal weight for sections for simplicity

    total_possible_score = sum(len(challenges) * rating_scores["Compliant"] * section_weights.get(section_item, 1) for section_item, challenges in validation_challenges.items())
    assessment_results["total_possible_score"] = total_possible_score

    for section_item, challenges in validation_challenges.items(): # Use the global validation_challenges dictionary
        assessment_results["sections"][section_item] = {
            "score": 0,
            "possible_score": len(challenges) * rating_scores["Compliant"],
            "status": "Not Assessed",
            "challenges": []
        }

        for challenge in challenges:
            prompt = f"""
You are a senior model risk validator assessing an IFRS 9 model documentation and code against a specific validation challenge.
Analyze the provided documentation and code to determine if the challenge is addressed.

Documentation Snippet:
{doc_text[:3000]}...

Model Code Snippet:
{code_content[:3000]}...

Validation Challenge: {challenge}

Instructions:
1. Rate how well the documentation and code address this challenge: 'Compliant', 'Partially Compliant', or 'Non-Compliant'.
2. Provide a concise rationale (1-3 sentences).
3. Identify any residual risk related to this challenge if it's not fully compliant.
4. Suggest potential mitigation actions.

Output format:
Rating: [Rating]
Rationale: [Rationale]
Residual Risk: [Residual Risk]
Mitigation: [Mitigation]
"""
            try:
                if not GEMINI_API_KEY:
                    raise ValueError("Gemini API key not configured.")
                # Use RAG if enabled to get additional context for the LLM
                rag_context_text = ""
                if rag_system:
                    rag_query = f"Validation challenge: {challenge}. IFRS 9 {st.session_state.get('model_type_sidebar', 'Model')} model validation."
                    retrieved_context = rag_system.retrieve_context(rag_query, k=st.session_state.get('retrieval_k', 3)) # Use fewer chunks for validation prompt
                    rag_context_text = "\nAdditional Context from Knowledge Base:\n" + "\n---\n".join([item['text'] for item in retrieved_context])

                full_prompt = prompt + rag_context_text

                gemini_model_name = st.session_state.get('ai_model', 'gemini-pro')
                # Default to a known working model if the selection is an old OpenAI model or potentially problematic
                if "gpt" in gemini_model_name.lower() or gemini_model_name == "gemini-pro": # Check for "gemini-pro" specifically
                    gemini_model_name = "gemini-1.5-flash-latest" # Or try "gemini-1.0-pro"
                model_instance = genai.GenerativeModel(gemini_model_name)
                generation_config = genai.types.GenerationConfig(
                    max_output_tokens=st.session_state.get('max_tokens', 500),
                    temperature=st.session_state.get('temperature', 0.3)
                )
                llm_response = model_instance.generate_content(full_prompt, generation_config=generation_config)
                output = llm_response.text
                lines = output.split("\n")
                rating = "Error" # Default
                rationale = "Could not parse rationale."
                residual_risk = "Could not parse residual risk."
                mitigation = "Could not parse mitigation."

                for line in lines:
                    if line.lower().startswith("rating:"):
                        rating = line.split(":", 1)[1].strip()
                    elif line.lower().startswith("rationale:"):
                        rationale = line.split(":", 1)[1].strip()
                    elif line.lower().startswith("residual risk:"):
                        residual_risk = line.split(":", 1)[1].strip()
                    elif line.lower().startswith("mitigation:"):
                        mitigation = line.split(":", 1)[1].strip()

                # Ensure rating is one of the expected values
                if rating not in rating_scores:
                    rating = "Error"
                    rationale = f"LLM returned unexpected rating format. Raw output: {output}"

                assessment_results["sections"][section_item]["challenges"].append({
                    "challenge": challenge,
                    "rating": rating,
                    "rationale": rationale,
                    "residual_risk": residual_risk,
                    "mitigation": mitigation
                })

                # Update counts and scores
                # Check if the key exists before incrementing
                count_key = f"{rating.lower().replace(' ', '_')}_count"
                if count_key in assessment_results:
                    assessment_results[count_key] += 1
                else: # If key doesn't exist (e.g. for "Error_count")
                    assessment_results["error_count"] += 1

                score_earned = rating_scores.get(rating, 0)
                assessment_results["sections"][section_item]["score"] += score_earned
                assessment_results["overall_score"] += score_earned * section_weights.get(section_item, 1)

                if rating != "Compliant":
                     assessment_results["gap_table"].append((section_item, challenge, rating, rationale, residual_risk, mitigation))

            except Exception as e:
                assessment_results["error_count"] += 1
                assessment_results["sections"][section_item]["challenges"].append({
                    "challenge": challenge,
                    "rating": "Error",
                    "rationale": f"API Error: {str(e)}. Ensure Gemini API key and model name are correct.",
                    "residual_risk": "N/A",
                    "mitigation": "N/A"
                })
                assessment_results["gap_table"].append((section_item, challenge, "Error", f"API Error: {str(e)}", "N/A", "N/A"))

        # Determine section status based on challenges
        section_challenges = assessment_results["sections"][section_item]["challenges"]
        if any(c['rating'] == 'Non-Compliant' for c in section_challenges):
            assessment_results["sections"][section_item]["status"] = "Non-Compliant"
        elif any(c['rating'] == 'Partially Compliant' for c in section_challenges):
             assessment_results["sections"][section_item]["status"] = "Partially Compliant"
        elif any(c['rating'] == 'Error' for c in section_challenges):
             assessment_results["sections"][section_item]["status"] = "Assessment Error"
        elif all(c['rating'] == 'Compliant' for c in section_challenges) and section_challenges: # Ensure there are challenges
             assessment_results["sections"][section_item]["status"] = "Compliant"
        else: # Default or if no challenges assessed
             assessment_results["sections"][section_item]["status"] = "Not Fully Assessed"


    # Determine overall risk rating based on overall score
    score_percentage = (assessment_results["overall_score"] / assessment_results["total_possible_score"]) * 100 if assessment_results["total_possible_score"] > 0 else 0

    if score_percentage >= 85:
        assessment_results["risk_rating"] = "A (Lowest Risk)"
    elif score_percentage >= 70:
        assessment_results["risk_rating"] = "B (Low Risk)"
    elif score_percentage >= 50:
        assessment_results["risk_rating"] = "C (Moderate Risk)"
    else:
        assessment_results["risk_rating"] = "D (Highest Risk)"

    assessment_results["summary"] = f"Overall Score: {assessment_results['overall_score']}/{assessment_results['total_possible_score']} ({score_percentage:.1f}%).\nRisk Rating: {assessment_results['risk_rating']}.\nCompliant: {assessment_results['compliant_count']}, Partially Compliant: {assessment_results['partially_compliant_count']}, Non-Compliant: {assessment_results['non_compliant_count']}, Errors: {assessment_results['error_count']}."

    return assessment_results

def display_rag_validation_results(results: Dict, validation_level: str):
    """
    Displays the RAG-enhanced validation results, including scoring and risk rating.
    """
    # Imports moved to top level
    try:
        from docx import Document as DocxDocumentForReport # Alias for report generation
        from docx.shared import Inches as DocxInches # For report generation
    except ImportError:
        DocxDocumentForReport = None
        st.warning("`python-docx` library not found. Report generation will be limited. Install with `pip install python-docx`.")

    st.subheader("📋 RAG-Enhanced Validation Results")

    # Overall Summary
    st.markdown(f"**Assessment Level:** `{validation_level}`")
    st.markdown(f"**Overall Score:** `{results.get('overall_score', 0)} / {results.get('total_possible_score', 0)}`")
    st.markdown(f"**Model Risk Rating:** `{results.get('risk_rating', 'N/A')}`")

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Compliant", results.get('compliant_count', 0))
    col2.metric("Partially Compliant", results.get('partially_compliant_count', 0))
    col3.metric("Non-Compliant", results.get('non_compliant_count', 0))
    col4.metric("Assessment Errors", results.get('error_count', 0))

    st.markdown("---")

    # Section Breakdown
    st.subheader("Detailed Section Assessment")
    for section_item, details in results.get('sections', {}).items():
        with st.expander(f"{section_item} - {details.get('status', 'N/A')} (Score: {details.get('score',0)}/{details.get('possible_score',0)})"):
            for challenge_result in details.get('challenges', []):
                st.markdown(f"**Challenge:** {challenge_result['challenge']}")
                st.markdown(f"**Rating:** `{challenge_result['rating']}`")
                st.markdown(f"**Rationale:** {challenge_result['rationale']}")
                if challenge_result.get('residual_risk') and challenge_result['residual_risk'] not in ["Could not parse residual risk.", "N/A"]:
                     st.markdown(f"**Residual Risk:** {challenge_result['residual_risk']}")
                if challenge_result.get('mitigation') and challenge_result['mitigation'] not in ["Could not parse mitigation.", "N/A"]:
                     st.markdown(f"**Mitigation:** {challenge_result['mitigation']}")
                st.markdown("---")

    # Gap Table
    if results.get('gap_table'):
        st.subheader("Validation Gap Summary")
        gap_df = pd.DataFrame(results['gap_table'], columns=['Section', 'Challenge', 'Rating', 'Rationale', 'Residual Risk', 'Mitigation'])
        st.dataframe(gap_df)

    # Generate Validation Report Button
    if st.button("📊 Generate Full Validation Report"):
        if DocxDocumentForReport:
            validation_doc = DocxDocumentForReport()
            validation_doc.add_heading("Model Validation Assessment Report", 0)
            validation_doc.add_paragraph(f"Assessment Date: {datetime.now().strftime('%B %d, %Y')}")
            validation_doc.add_paragraph(f"Assessment Level: {validation_level}")
            validation_doc.add_paragraph(f"Overall Score: {results.get('overall_score', 0)} / {results.get('total_possible_score', 0)}")
            validation_doc.add_paragraph(f"Model Risk Rating: {results.get('risk_rating', 'N/A')}")
            validation_doc.add_paragraph(f"Summary: Compliant: {results.get('compliant_count', 0)}, Partially Compliant: {results.get('partially_compliant_count', 0)}, Non-Compliant: {results.get('non_compliant_count', 0)}, Errors: {results.get('error_count', 0)}")

            for section_item, details in results.get('sections', {}).items():
                validation_doc.add_heading(section_item, level=1)
                validation_doc.add_paragraph(f"Status: {details.get('status','N/A')} (Score: {details.get('score',0)}/{details.get('possible_score',0)})")

                for challenge_result in details.get('challenges',[]):
                    validation_doc.add_paragraph(f"Challenge: {challenge_result['challenge']}", style='ListBullet')
                    validation_doc.add_paragraph(f"Rating: {challenge_result['rating']}")
                    validation_doc.add_paragraph(f"Rationale: {challenge_result['rationale']}")
                    if challenge_result.get('residual_risk') and challenge_result['residual_risk'] not in ["Could not parse residual risk.", "N/A"]:
                        validation_doc.add_paragraph(f"Residual Risk: {challenge_result['residual_risk']}")
                    if challenge_result.get('mitigation') and challenge_result['mitigation'] not in ["Could not parse mitigation.", "N/A"]:
                        validation_doc.add_paragraph(f"Mitigation: {challenge_result['mitigation']}")
                    validation_doc.add_paragraph("") # Add space

            if results.get('gap_table'):
                validation_doc.add_heading("Validation Gap Summary", level=1)
                table = validation_doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Section'
                hdr_cells[1].text = 'Challenge'
                hdr_cells[2].text = 'Rating'
                hdr_cells[3].text = 'Rationale'
                hdr_cells[4].text = 'Residual Risk'
                hdr_cells[5].text = 'Mitigation'

                for row_data in results['gap_table']:
                    cells = table.add_row().cells
                    for i_cell, cell_text in enumerate(row_data):
                        cells[i_cell].text = str(cell_text)[:250] # Truncate long text for table cells

            val_doc_buffer = BytesIO()
            validation_doc.save(val_doc_buffer)
            val_doc_buffer.seek(0)
            
            val_report_filename = f"ifrs9_validation_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            st.download_button(
                label="📥 Download Validation Report",
                data=val_doc_buffer,
                file_name=val_report_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Could not generate DOCX report. `python-docx` might be missing.")

def get_knowledge_base_stats(rag_system: SimpleRAG) -> Dict:
    """Provides statistics about the knowledge base."""
    stats = {
        'total_docs_in_store': len(rag_system.vector_store.documents), # This is actually total chunks
        'total_chunks': len(rag_system.vector_store.documents), 
        'vocab_size': len(rag_system.embedding_model.vocabulary) if hasattr(rag_system.embedding_model, 'vocabulary') else 'N/A (Custom Embedding)',
        'avg_chunk_length': 0, # Approximate
        'doc_types': defaultdict(int)
    }
    
    total_words = 0
    if stats['total_chunks'] > 0:
        for doc_text in rag_system.vector_store.documents: # Iterate over chunk texts
            total_words += len(doc_text.split())
        stats['avg_chunk_length'] = round(total_words / stats['total_chunks']) if stats['total_chunks'] > 0 else 0
    
    for meta in rag_system.vector_store.metadata: # Iterate over metadata of chunks
        stats['doc_types'][meta.get('type', 'unknown')] += 1
        
    return stats

def display_rag_workflow(workflow_steps: List[Dict]):
    """Displays the RAG workflow steps in the UI."""
    for step in workflow_steps:
        status_icon = {
            'info': '🔵',
            'processing': '🟡',
            'success': '🟢',
            'error': '🔴'
        }.get(step['status'], '⚪')
        
        st.markdown(f"{status_icon} **{step['timestamp']}** - {step['step']}")
        if step.get('details'):
            st.markdown(f"  > *{step['details']}*")

# =============================================
# STREAMLIT APPLICATION
# =============================================

# Initialize session state
if 'rag_system' not in st.session_state:
    st.session_state.rag_system = SimpleRAG()
if 'knowledge_base_built' not in st.session_state:
    st.session_state.knowledge_base_built = False
if 'generated_doc_buffer' not in st.session_state: # Store the BytesIO buffer
    st.session_state.generated_doc_buffer = None
if 'generated_doc_filename' not in st.session_state: # Store the filename
    st.session_state.generated_doc_filename = None
if 'validation_results' not in st.session_state:
    st.session_state.validation_results = None
if 'code_analysis' not in st.session_state:
    st.session_state.code_analysis = None
if 'code_content_for_validation' not in st.session_state: # Store code content for validation tab
    st.session_state.code_content_for_validation = None
if 'ai_model' not in st.session_state:
    st.session_state.ai_model = "gpt-4" # Default AI model
if 'temperature' not in st.session_state:
    st.session_state.temperature = 0.3 # Default temperature
if 'max_tokens' not in st.session_state:
    st.session_state.max_tokens = 1500 # Default max_tokens
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'user_info' not in st.session_state:
    st.session_state.user_info = None
if 'auth_view' not in st.session_state: # To switch between login/signup
    st.session_state.auth_view = "login"

# Header
st.title("🧠 AI-Powered Risk Model Documentation with RAG Intelligence")
st.markdown("**Effortlessly generate model documentation using cutting-edge AI.**")

# Sidebar Configuration
with st.sidebar:
    st.markdown("## ⚙️ Configuration")
    
    # RAG Settings
    st.markdown("### 🧠 RAG Configuration")
    # Store these in session state so they persist across runs and can be used by RAG system
    st.session_state.chunk_size = st.slider("Chunk Size (words)", 200, 800, st.session_state.get('chunk_size', 400), key="chunk_size_slider")
    st.session_state.overlap = st.slider("Chunk Overlap (words)", 20, 100, st.session_state.get('overlap', 50), key="overlap_slider")
    st.session_state.retrieval_k = st.slider("Retrieval Top-K", 3, 10, st.session_state.get('retrieval_k', 5), key="retrieval_k_slider")
    
    # Model Configuration for Documentation Generation
    st.markdown("### 🎯 Model Settings (for Documentation)")
    st.session_state.model_type_sidebar = st.selectbox("Model Type", ["PD (Probability of Default)", "LGD (Loss Given Default)", "EAD (Exposure at Default)"], key="doc_model_type")
    st.session_state.portfolio_type_sidebar = st.selectbox("Portfolio Type", ["Credit Cards", "Personal Loans", "Mortgages", "Corporate Lending", "SME"], key="doc_portfolio_type")
    
    # Regulatory Framework for Documentation Generation
    st.markdown("### 📋 Regulatory Framework (for Documentation)")
    st.session_state.regulations_sidebar = st.multiselect(
        "Applicable Regulations", 
        ["IFRS 9", "PRA SS1/23", "PRA SS11/13", "TRIM", "Basel III", "CRR II"],
        default=["IFRS 9"],
        key="doc_regulations"
    )
    
    st.markdown("---")
    st.markdown("**AlgoRisk AI Platform**")
    st.markdown("*RAG-Powered Risk Intelligence*")

# --- Authentication Functions ---
def signup_user(name, email, password):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    try:
        hashed_password = generate_password_hash(password, method='pbkdf2:sha256')
        c.execute("INSERT INTO users (name, email, password_hash) VALUES (?, ?, ?)", (name, email, hashed_password))
        conn.commit()
        return True
    except sqlite3.IntegrityError: # Email already exists
        return False
    finally:
        conn.close()

def verify_login(email, password):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT name, email, password_hash FROM users WHERE email = ?", (email,))
    user_record = c.fetchone()
    conn.close()
    if user_record and check_password_hash(user_record[2], password):
        return {"name": user_record[0], "email": user_record[1]}
    return None

def render_auth_page():
    st.markdown("<div class='auth-container'>", unsafe_allow_html=True) # Apply auth-specific styles

    if st.session_state.auth_view == "login":
        st.header("Login to AlgoRisk AI")
        with st.form("login_form"):
            email = st.text_input("Email")
            password = st.text_input("Password", type="password")
            submit_login = st.form_submit_button("Login")

            if submit_login:
                if not email or not password:
                    st.error("Please enter both email and password.")
                else:
                    user = verify_login(email, password)
                    if user:
                        st.session_state.logged_in = True
                        st.session_state.user_info = user
                        st.rerun()
                    else:
                        st.error("Invalid email or password.")
        
        if st.button("Login with Google (Placeholder)"):
            st.info("""
            **Google OAuth Integration:**
            This would typically redirect you to Google for authentication.
            Upon successful Google login, the app would receive your profile information.
            Consider using libraries like `streamlit-oauth` for easier integration.
            For now, please use email/password.
            """)

        st.markdown("---")
        if st.button("Don't have an account? Sign Up"):
            st.session_state.auth_view = "signup"
            st.rerun()

    elif st.session_state.auth_view == "signup":
        st.header("Create New Account")
        with st.form("signup_form"):
            name = st.text_input("Full Name")
            email = st.text_input("Email")
            password = st.text_input("Password", type="password")
            confirm_password = st.text_input("Confirm Password", type="password")
            submit_signup = st.form_submit_button("Sign Up")

            if submit_signup:
                if not name or not email or not password or not confirm_password:
                    st.error("Please fill all fields.")
                elif password != confirm_password:
                    st.error("Passwords do not match.")
                elif signup_user(name, email, password):
                    st.success("Account created successfully! Please login.")
                    st.session_state.auth_view = "login" # Switch to login view
                    st.rerun()
                else:
                    st.error("Email already exists or another error occurred.")
        
        st.markdown("---")
        if st.button("Already have an account? Login"):
            st.session_state.auth_view = "login"
            st.rerun()
    
    st.markdown("</div>", unsafe_allow_html=True)

# --- Main App Logic ---
if not st.session_state.logged_in:
    # Clear out main app content if not logged in, show auth page
    # The st.set_page_config is global, so title/icon remain.
    # We can hide the main app's title if we want a cleaner auth screen.
    # For now, the auth page will appear within the main app's structure.
    render_auth_page()
    st.stop() # Stop further execution of the main app

# If logged in, show a welcome message and logout button in the sidebar
st.sidebar.markdown("---")
st.sidebar.subheader(f"Welcome, {st.session_state.user_info['name']}!")
if st.sidebar.button("Logout"):
    st.session_state.logged_in = False
    st.session_state.user_info = None
    st.session_state.auth_view = "login" # Reset to login view on logout
    st.rerun()


tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🧠 RAG Knowledge Base",         # Great already – clear and smart
    "📄 Auto-Generated Documentation",  # More descriptive than "Model Documentation"
    "✅ Model Validation & Compliance", # Adds clarity and governance feel
    "💬 Interactive RAG Assistant",     # Modern and engaging
    "📊 Insights & Performance Analytics"  # Richer and more specific
])

with tab1:
    st.header("🧠 RAG Knowledge Base Management")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("📚 Build Knowledge Base")
        
        # File uploaders for knowledge base
        st.markdown("**Upload Reference Documents:**")
        regulatory_docs = st.file_uploader(
            "Regulatory Guidelines (PDF, TXT, DOCX)", 
            type=["pdf", "txt", "docx"], 
            accept_multiple_files=True,
            key="reg_docs"
        )
        
        model_docs = st.file_uploader(
            "Model Documentation Templates", 
            type=["pdf", "txt", "docx"], 
            accept_multiple_files=True,
            key="model_docs"
        )
        
        code_files = st.file_uploader(
            "Reference Code Files (.py)", 
            type=["py"], 
            accept_multiple_files=True,
            key="code_files"
        )
        
        # Include built-in knowledge
        include_builtin = st.checkbox("Include Built-in IFRS 9 Knowledge Base", True)
        
        if st.button("🔨 Build RAG Knowledge Base", type="primary"):
            documents = []
            
            # Add built-in knowledge
            if include_builtin:
                builtin_docs = get_builtin_ifrs9_knowledge()
                documents.extend(builtin_docs)
            
            # Process uploaded files
            if regulatory_docs:
                for file in regulatory_docs:
                    content = extract_text_from_file(file)
                    if content and not content.startswith("Error:"): # Check for successful extraction
                        documents.append({
                            'content': content,
                            'source': file.name,
                            'type': 'regulatory'
                        })
                    elif content:
                         st.warning(f"Skipping {file.name} due to extraction error: {content}")
            
            if model_docs:
                for file in model_docs:
                    content = extract_text_from_file(file)
                    if content and not content.startswith("Error:"): # Check for successful extraction
                        documents.append({
                            'content': content,
                            'source': file.name,
                            'type': 'model_doc'
                        })
                    elif content:
                         st.warning(f"Skipping {file.name} due to extraction error: {content}")
            
            if code_files:
                for file in code_files:
                    try:
                        file.seek(0) # Reset file pointer
                        content = file.read().decode('utf-8', errors='ignore')
                        documents.append({
                            'content': content,
                            'source': file.name,
                            'type': 'code'
                        })
                    except Exception as e:
                        st.warning(f"Skipping {file.name} due to read error: {e}")
            
            if documents:
                # Ensure RAG system uses current sidebar settings for chunking
                st.session_state.rag_system.chunker = DocumentChunker(
                    chunk_size=st.session_state.chunk_size, 
                    overlap=st.session_state.overlap
                )
                
                with st.spinner("Building RAG knowledge base..."):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    def progress_callback(progress, message):
                        progress_bar.progress(progress / 100) # Progress bar expects 0-1
                        status_text.text(message)
                    
                    try:
                        num_chunks = st.session_state.rag_system.build_knowledge_base(documents, progress_callback)
                        st.session_state.knowledge_base_built = True
                        st.success(f"✅ Knowledge base built successfully with {num_chunks} chunks from {len(documents)} documents!")
                    except Exception as e:
                         st.error(f"Error building knowledge base: {e}")
                         st.session_state.knowledge_base_built = False # Mark as not built on error

            else:
                st.error("Please upload at least one document or enable built-in knowledge base.")
    
    with col2:
        st.subheader("📊 Knowledge Base Status")
        if st.session_state.knowledge_base_built:
            st.success("✅ Knowledge Base Ready")
            kb_stats = get_knowledge_base_stats(st.session_state.rag_system)
            st.metric("Total Chunks", kb_stats.get('total_chunks', 0)) # Display chunks as they are the indexed units
            st.metric("Vocabulary Size", kb_stats.get('vocab_size', 'N/A'))
            st.metric("Average Chunk Length", f"{kb_stats.get('avg_chunk_length', 0)} words")
            
            # Document type breakdown
            if kb_stats.get('doc_types'):
                st.subheader("📑 Document Types")
                for doc_type, count in kb_stats['doc_types'].items():
                    st.write(f"• {doc_type.title()}: {count} documents")
        else:
            st.info("🔄 Knowledge Base Not Built")
        
        # Display recent workflow steps
        if st.session_state.rag_system.workflow_steps:
            st.subheader("🔄 Recent Activity")
            display_rag_workflow(st.session_state.rag_system.workflow_steps[-5:]) # Show last 5 steps

with tab2:
    st.header("📤 Model Upload & AI Documentation Generation")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📁 Upload Model Files")
        code_file = st.file_uploader("Upload Model Code (.py)", type=["py"], key="main_code")
        data_file = st.file_uploader("Upload Sample Data (.csv)", type=["csv"], key="data_file")
        config_file = st.file_uploader("Upload Configuration (.json)", type=["json"], key="config_file")
    
    # Get sidebar selections for documentation generation
    selected_model_type = st.session_state.get('doc_model_type', "PD (Probability of Default)")
    selected_portfolio_type = st.session_state.get('doc_portfolio_type', "Credit Cards")
    selected_regulations = st.session_state.get('doc_regulations', ["IFRS 9"])

    with col2:
        st.subheader("📈 Model Results")
        results_input = st.text_area(
            "Model Performance Metrics (JSON format)",
            placeholder='{"auc": 0.75, "gini": 0.50, "ks": 0.35, "psi": 0.12}',
            height=100
        )
        
        additional_info = st.text_area(
            "Additional Context",
            placeholder="Any additional information about the model, assumptions, or special considerations...",
            height=100
        )
    
    # Enhanced RAG-powered generation
    use_rag = st.checkbox("🧠 Use RAG for Enhanced Documentation", value=True)
    
    # Generate Documentation Button
    if st.button("🚀 Generate RAG-Enhanced Documentation", type="primary"):
        if not GEMINI_API_KEY:
             st.error("Gemini API key not configured. Please set the GEMINI_API_KEY environment variable.")
        elif code_file:
            with st.spinner("🤖 AI is analyzing your model with RAG enhancement..."):
                code_file.seek(0) # Reset file pointer before reading
                code_content = code_file.read().decode("utf-8", errors='ignore')
                
                # Show workflow
                workflow_container = st.container()
                with workflow_container:
                    st.subheader("🔄 RAG-Enhanced Generation Workflow")
                    workflow_placeholder = st.empty() # Placeholder for progress/status
                
                # Analyze code
                workflow_placeholder.text("Analyzing model code structure...")
                st.session_state.code_analysis = analyze_code_structure(code_content)
                st.session_state.code_content_for_validation = code_content # Store for validation tab
                
                # Parse results
                workflow_placeholder.text("Processing model results...")
                try:
                    results = json.loads(results_input) if results_input else {}
                except Exception as e:
                    results = {"raw_results": results_input, "parse_error": str(e)}
                    st.warning(f"Could not parse results JSON: {e}. Using raw input.")
                
                # Read data and config files (simple text read for now)
                data_content_str = ""
                if data_file:
                    data_file.seek(0) # Reset file pointer
                    data_content_str = data_file.read().decode("utf-8", errors='ignore')
                config_content_str = ""
                if config_file:
                    config_file.seek(0) # Reset file pointer
                    config_content_str = config_file.read().decode("utf-8", errors='ignore')

                # RAG-enhanced generation
                if use_rag and st.session_state.knowledge_base_built:
                    try:
                        doc_content = generate_rag_enhanced_documentation(
                            st.session_state.rag_system,
                            st.session_state.code_analysis,
                            results,
                            selected_model_type,
                            selected_portfolio_type,
                            selected_regulations,
                            code_content,
                            data_content_str,
                            config_content_str,
                            additional_info,
                            workflow_placeholder # Pass the placeholder
                        )
                    except Exception as e:
                        st.error(f"RAG-enhanced documentation generation failed: {str(e)}. Falling back to standard generation.")
                        # Ensure fallback function has correct signature
                        doc_content = generate_documentation_content(
                            st.session_state.code_analysis, 
                            results, 
                            selected_model_type,
                            selected_portfolio_type,
                            selected_regulations,
                            code_content_str=code_content, # pass code_content
                            data_content_str=data_content_str, # pass data_content
                            config_content_str=config_content_str # pass config_content
                        )
                else:
                    # Standard generation without RAG
                    doc_content = generate_documentation_content(
                        st.session_state.code_analysis, 
                        results, 
                        selected_model_type, 
                        selected_portfolio_type,
                        selected_regulations,
                        code_content_str=code_content, # pass code_content
                        data_content_str=data_content_str, # pass data_content
                        config_content_str=config_content_str # pass config_content
                    )
                
                # Create Word document
                workflow_placeholder.text("Creating Word document...")
                doc_buffer, doc_filename = create_word_document(doc_content, selected_model_type)
                st.session_state.generated_doc_buffer = doc_buffer # Store buffer in state
                st.session_state.generated_doc_filename = doc_filename # Store filename
                
                workflow_placeholder.text("Documentation generated successfully!")
                st.success("✅ Documentation generated successfully!")
                
                # Download button using the buffer and filename from create_word_document
                st.download_button(
                    label="📥 Download Complete Documentation",
                    data=st.session_state.generated_doc_buffer,
                    file_name=st.session_state.generated_doc_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("Please upload at least the model code file to proceed.")

with tab3:
    st.header("🔍 Model Validation Dashboard")
    
    # Check if code analysis (and thus code) is available from Tab 2
    if st.session_state.get('code_analysis') and st.session_state.get('code_content_for_validation'): 
        st.subheader("Upload Documentation for Validation")
        validation_doc_file = st.file_uploader("Upload Documentation to Validate (.docx, .txt)", type=["docx", "txt"], key="validation_doc_file")

        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.subheader("🎯 RAG-Enhanced Validation")
            
            validation_level = st.selectbox("Assessment Depth", ["Quick Check", "Standard Review", "Comprehensive Audit"])
            use_rag_validation = st.checkbox("🧠 Use RAG for Validation Enhancement", value=True)

            if st.button("🔍 Run RAG-Enhanced Validation"):
                if not GEMINI_API_KEY:
                    st.error("Gemini API key not configured. Please set the GEMINI_API_KEY environment variable.")
                elif validation_doc_file:
                    with st.spinner("Running RAG-enhanced validation..."):
                        validation_doc_text = extract_text_from_file(validation_doc_file)
                        code_content_for_val = st.session_state.code_content_for_validation # Get code content from session state
                        
                        if validation_doc_text and not validation_doc_text.startswith("Error:"):
                            validation_results = run_rag_validation_assessment(
                                st.session_state.rag_system if use_rag_validation and st.session_state.knowledge_base_built else None,
                                st.session_state.code_analysis, # Pass code analysis
                                validation_level,
                                validation_doc_text, 
                                code_content_for_val # Pass code text
                            )
                            st.session_state.validation_results = validation_results
                        elif validation_doc_text:
                             st.error(f"Error reading validation document: {validation_doc_text}")
                        else:
                             st.error("Could not read content from the uploaded validation document.")
                else:
                    st.warning("Please upload the documentation file to validate.")
        
        with col2:
            st.subheader("📊 Validation Metrics")
            # Display quick metrics if validation results are available
            if st.session_state.validation_results:
                results = st.session_state.validation_results
                # Use .get() for safety
                st.metric("Overall Score", f"{results.get('overall_score', 0)}/{results.get('total_possible_score', 0)}")
                st.metric("Compliant Items", results.get('compliant_count', 0))
                st.metric("Issues Found", results.get('non_compliant_count', 0) + results.get('partially_compliant_count', 0) + results.get('error_count', 0))
        
        # Display detailed validation results if available and a document was uploaded for validation
        if st.session_state.validation_results and validation_doc_file: 
            display_rag_validation_results(st.session_state.validation_results, validation_level)
        elif st.session_state.get('code_analysis'):
             st.info("Upload the documentation file you wish to validate and click 'Run RAG-Enhanced Validation'.")
    else:
        st.info("Please upload model code in the 'Upload & Generate' tab first to access validation features.")

with tab4:
    st.header("💬 RAG-Powered Chat Interface")
    
    if not GEMINI_API_KEY:
         st.error("Gemini API key not configured. Please set the GEMINI_API_KEY environment variable to use the chat.")
    elif st.session_state.knowledge_base_built:
        st.subheader("🤖 Ask Questions About Your Model & IFRS 9")
        
        # Chat interface
        query = st.text_input("Enter your question:", placeholder="e.g., What are the key requirements for IFRS 9 staging?")
        
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("🔍 Ask RAG", type="primary"):
                if query:
                    with st.spinner("Processing your question with RAG..."):
                        # Perform RAG query using retrieval_k from session state
                        rag_result = st.session_state.rag_system.query(query, k=st.session_state.get('retrieval_k', 5))
                        
                        # Display workflow
                        st.subheader("🔄 RAG Processing Workflow")
                        display_rag_workflow(rag_result['workflow'])
                        
                        # Display response
                        st.subheader("🤖 AI Response")
                        st.markdown(rag_result['response'])
                        
                        # Display context
                        if rag_result['context']:
                            with st.expander("📚 Retrieved Context"):
                                for i, item in enumerate(rag_result['context'], 1):
                                    st.markdown(f"**Context {i}** (Similarity: {item['similarity']:.1%})")
                                    st.markdown(f"*Source: {item['metadata'].get('source', 'Unknown')}*")
                                    with st.container():
                                        st.markdown(f'<div class="chunk-preview">{item["text"][:300]}...</div>', unsafe_allow_html=True)
                else:
                    st.error("Please enter a question.")
        
        with col2:
            st.markdown("**Example Questions:**")
            example_questions = [
                "What are the key IFRS 9 staging criteria?",
                "How should I validate a PD model?",
                "What documentation is required for model governance?",
                "How do I implement significant increase in credit risk triggers?",
                "What are common validation challenges for IFRS 9 models?"
            ]
            
            # Use session state to trigger example query
            for question in example_questions:
                if st.button(f"💡 {question}", key=f"example_{hash(question)}"):
                    # Set the query input field value and trigger the button click
                    # This requires a bit more advanced Streamlit state management or JS injection
                    # For simplicity, we'll just display the query here.
                    st.info(f"You selected: '{question}'. Please paste this into the input box above and click 'Ask RAG'.")
                    # A more complex approach would involve setting st.session_state.query_input = question
                    # and potentially rerunning the app or using callbacks.
    else:
        st.info("Please build the RAG knowledge base first to use the chat interface.")

with tab5:
    st.header("📊 RAG Analytics & Model Insights")
    
    if st.session_state.knowledge_base_built or st.session_state.get('code_analysis'):
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("🧠 Knowledge Base Analytics")
            if st.session_state.knowledge_base_built:
                kb_stats = get_knowledge_base_stats(st.session_state.rag_system)
                
                st.metric("Total Chunks", kb_stats.get('total_chunks', 0))
                st.metric("Vocabulary Size", kb_stats.get('vocab_size', 'N/A'))
                st.metric("Average Chunk Length", f"{kb_stats.get('avg_chunk_length', 0)} words")
                
                # Document type breakdown
                if kb_stats.get('doc_types'):
                    st.subheader("📑 Document Types")
                    for doc_type, count in kb_stats['doc_types'].items():
                        st.write(f"• {doc_type.title()}: {count} documents")
            else:
                 st.info("Build the knowledge base to see analytics.")

        with col2:
            st.subheader("🔧 Model Code Analysis")
            if st.session_state.code_analysis:
                complexity = st.session_state.code_analysis.get('complexity', {})
                
                st.metric("Lines of Code", complexity.get('lines', 0))
                st.metric("Functions", complexity.get('functions', 0))
                st.metric("Classes", complexity.get('classes', 0))
                st.metric("Imports", complexity.get('imports', 0))
                
                # Top libraries
                libraries = st.session_state.code_analysis.get('libraries', [])
                if libraries:
                    st.subheader("📚 Key Libraries")
                    for lib in libraries[:5]:
                        lib_name = lib # Libraries list already contains names
                        if lib_name:
                            st.write(f"• {lib_name}")
                
                # Identified Features
                features = st.session_state.code_analysis.get('features', [])
                if features:
                    st.subheader("🎯 Identified Features")
                    for feature in features[:10]: # Show top 10 identified features
                        st.write(f"• {feature}")

            else:
                st.info("Upload and analyze a model code file to see analytics.")

# Footer
st.markdown("---")
st.markdown("### 🔗 Quick Links")
col1, col2, col3, col4 = st.columns(4)
with col1:
    st.markdown("📚 Documentation")
with col2:
    st.markdown("🆘 Support")
with col3:
    st.markdown("📊 Examples")
with col4:
    st.markdown("🔧 API Access")

st.markdown("**© 2025 AlgoRisk AI - Transforming Risk Management with Artificial Intelligence**")

# === FastAPI Integration for Streamlit app ===
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import os
import google.generativeai as genai
from typing import List
import uvicorn
import tempfile
from docx import Document

# --- FastAPI App ---
fastapi_app = FastAPI()

fastapi_app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Gemini Setup ---
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
gemini_model = genai.GenerativeModel("gemini-pro")

# --- Schemas ---
class ValidationRequest(BaseModel):
    code: str
    level: str

class ChatRequest(BaseModel):
    message: str
    history: List[str] = []

class BuildKBRequest(BaseModel):
    documents: List[str]

# === /run-validation ===
@fastapi_app.post("/run-validation")
def run_validation(request: ValidationRequest):
    return {
        "overallScore": 85,
        "riskRating": "B",
        "compliant": 3,
        "nonCompliant": 1,
        "warnings": 1,
        "challenges": [
            {"category": "Data Quality", "status": "compliant", "score": 92, "details": "Checks passed."},
            {"category": "Model Performance", "status": "compliant", "score": 88, "details": "Performance OK."},
            {"category": "Regulatory Compliance", "status": "warning", "score": 75, "details": "IFRS 9 gaps."},
            {"category": "Code Quality", "status": "non-compliant", "score": 65, "details": "Poor coverage."},
            {"category": "Risk Management", "status": "compliant", "score": 90, "details": "Good framework."}
        ]
    }

# === /generate-documentation ===
@fastapi_app.post("/generate-documentation")
def generate_documentation(code: str = Form(...), comments: str = Form("")):
    prompt = f"Generate full model documentation for the following code:\n\n{code}\n\nAdditional comments: {comments}"
    response = gemini_model.generate_content(prompt)
    return {"documentation": response.text}

# === /chat-rag ===
@fastapi_app.post("/chat-rag")
def chat_rag(request: ChatRequest):
    history = "\n".join(request.history)
    prompt = f"{history}\nUser: {request.message}\nAssistant:"
    response = gemini_model.generate_content(prompt)
    return {"reply": response.text}

# === /build-knowledge-base ===
@fastapi_app.post("/build-knowledge-base")
def build_knowledge_base(payload: BuildKBRequest):
    return {"status": "success", "indexed_docs": len(payload.documents)}

# === /download-doc ===
@fastapi_app.get("/download-doc")
def download_document():
    doc = Document()
    doc.add_heading("Model Documentation", 0)
    doc.add_paragraph("This is a placeholder Word document generated by the backend.")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return FileResponse(tmp.name, filename="model_doc.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# === /download-validation-report ===
@fastapi_app.get("/download-validation-report")
def download_validation():
    doc = Document()
    doc.add_heading("Model Validation Report", 0)
    doc.add_paragraph("Sample validation results...")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return FileResponse(tmp.name, filename="validation_report.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# === Run FastAPI Locally ===

    from fastapi import FastAPI

fastapi_app = FastAPI()

# define endpoints...

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:fastapi_app", host="0.0.0.0", port=8000)

