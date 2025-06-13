from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import tempfile
import os
import shutil

from utils.rag_engine import SimpleRAG
from utils.doc_gen import generate_doc_from_inputs
from utils.validator import validate_model
from utils.doc_gen import create_word_doc
from docx import Document

app = FastAPI()

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load Gemini API
import google.generativeai as genai
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
else:
    raise Exception("Gemini API key not found. Set GEMINI_API_KEY as environment variable.")

# Shared RAG system
rag_system = SimpleRAG()

# === Endpoints ===

@app.post("/build-knowledge-base")
async def build_kb(files: list[UploadFile] = File(...)):
    docs = []
    for file in files:
        content = await file.read()
        text = content.decode("utf-8", errors="ignore")
        docs.append({'content': text, 'source': file.filename, 'type': 'uploaded'})
    chunk_count = rag_system.build_knowledge_base(docs)
    return {"status": "success", "indexed_chunks": chunk_count}

@app.post("/generate-documentation")
async def generate_doc(
    code: UploadFile = File(...),
    data: UploadFile = File(None),
    config: UploadFile = File(None),
    comments: str = Form(""),
    model_type: str = Form("PD"),
    portfolio_type: str = Form("Credit Card"),
    regulations: str = Form("IFRS 9"),
):
    code_content = await code.read()
    code_str = code_content.decode("utf-8", errors="ignore")

    data_str = await data.read() if data else b""
    config_str = await config.read() if config else b""

    rag_result = generate_doc_from_inputs(
        rag_system, code_str, data_str.decode("utf-8"), config_str.decode("utf-8"),
        model_type, portfolio_type, regulations.split(","), comments
    )

    # Generate Word
    doc_buffer, filename = create_word_doc(rag_result, model_type)
    temp_file_path = f"/tmp/{filename}"
    with open(temp_file_path, "wb") as f:
        f.write(doc_buffer.read())

    return FileResponse(temp_file_path, filename=filename)

@app.post("/run-validation")
async def validate(
    code: UploadFile = File(...),
    doc: UploadFile = File(...),
    level: str = Form("Standard")
):
    code_str = (await code.read()).decode("utf-8", errors="ignore")
    doc_str = (await doc.read()).decode("utf-8", errors="ignore")
    result = validate_model(rag_system, code_str, doc_str, level)
    return JSONResponse(content=result)

@app.post("/chat-rag")
async def chat_rag(request: Request):
    data = await request.json()
    message = data.get("message", "")
    response = rag_system.query(message)
    return {"response": response.get("response", ""), "workflow": response.get("workflow", [])}

@app.get("/download-doc")
def download_doc():
    doc = Document()
    doc.add_heading("Model Documentation", 0)
    doc.add_paragraph("This is a placeholder.")
    path = "/tmp/doc_placeholder.docx"
    doc.save(path)
    return FileResponse(path, filename="model_doc.docx")

@app.get("/download-validation-report")
def download_val():
    doc = Document()
    doc.add_heading("Validation Report", 0)
    doc.add_paragraph("Sample validation result.")
    path = "/tmp/validation_report.docx"
    doc.save(path)
    return FileResponse(path, filename="validation_report.docx")
