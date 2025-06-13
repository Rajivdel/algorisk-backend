import os
from typing import Dict, List, Tuple, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# --- Documentation Generation Utilities ---
def generate_documentation_content(code_analysis: Dict, results: Dict, model_type: str, portfolio_type: str, regulations: List[str], code_content_str: str = None, data_content_str: str = None, config_content_str: str = None) -> Dict[str, str]:
    """
    Generate documentation content sections as a dictionary.
    """
    doc_content = {}
    doc_content["Executive Summary"] = f"This document describes the {model_type} model for {portfolio_type} under {', '.join(regulations)}."
    doc_content["Model Purpose and Use"] = code_analysis.get("purpose", "Purpose not detected.")
    doc_content["Portfolio Segmentation"] = code_analysis.get("segmentation", "Segmentation not detected.")
    doc_content["Data Inputs and Sources"] = data_content_str or "Data inputs not provided."
    doc_content["Model Methodology"] = code_analysis.get("methodology", "Methodology not detected.")
    doc_content["Feature Engineering and Variable Selection"] = code_analysis.get("features", "Features not detected.")
    doc_content["Model Training and Validation Strategy"] = code_analysis.get("validation", "Validation strategy not detected.")
    doc_content["Model Performance Results"] = str(results) if results else "Results not provided."
    doc_content["Assumptions and Limitations"] = code_analysis.get("assumptions", "Assumptions not detected.")
    doc_content["Regulatory Compliance Assessment"] = f"Assessed against: {', '.join(regulations)}."
    # Add more sections as needed
    return doc_content

def create_word_document(doc_content: Dict[str, str], model_type: str) -> Tuple[BytesIO, str]:
    """
    Create a Word document from doc_content and return a BytesIO buffer and filename.
    """
    doc = Document()
    doc.add_heading(f"{model_type} Model Documentation", 0)
    for section, content in doc_content.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"{model_type.replace(' ', '_').lower()}_documentation.docx"
    return buffer, filename

# --- RAG-Enhanced Documentation Generation ---
def generate_rag_enhanced_documentation(rag_system, code_analysis: Dict, results: Dict, model_type: str, portfolio_type: str, regulations: List[str], code_content: str, data_content: str, config_content: str, additional_info: str, workflow_placeholder=None) -> Dict[str, str]:
    """
    Use the RAG system to enhance documentation content.
    """
    # Retrieve relevant context from RAG
    rag_context = rag_system.retrieve_context(f"{model_type} {portfolio_type} {','.join(regulations)}", k=5)
    context_text = '\n'.join([c.get('text', '') for c in rag_context])
    doc_content = generate_documentation_content(
        code_analysis,
        results,
        model_type,
        portfolio_type,
        regulations,
        code_content_str=code_content,
        data_content_str=data_content,
        config_content_str=config_content
    )
    # Optionally, add RAG context to each section
    for section in doc_content:
        doc_content[section] += f"\n\n[RAG Context]\n{context_text}"
    if additional_info:
        doc_content["Additional Information"] = additional_info
    if workflow_placeholder:
        workflow_placeholder.text("RAG-enhanced documentation generated.")
    return doc_content
import os
from docx import Document
from typing import List, Dict

def create_docx_documentation(sections: List[Dict[str, str]], save_path: str) -> str:
    """
    Generate a Word document from AI-generated documentation sections.

    Args:
        sections: List of sections with title and content
        save_path: File path where the document will be saved

    Returns:
        The path to the saved Word document
    """
    doc = Document()
    doc.add_heading("Model Documentation", 0)

    for section in sections:
        doc.add_heading(section["title"], level=1)
        doc.add_paragraph(section["content"])

    doc.save(save_path)
    return save_path


def simulate_ai_documentation(code_text: str, user_comments: str = "") -> List[Dict[str, str]]:
    """
    Simulates AI-generated model documentation from uploaded code and comments.

    Args:
        code_text: Source code content uploaded by user
        user_comments: Optional text area input from user

    Returns:
        A list of sections (each with title and content)
    """
    doc_sections = [
        "Executive Summary",
        "Model Overview",
        "Data Description",
        "Methodology",
        "Model Development",
        "Model Validation",
        "Performance Metrics",
        "Limitations",
        "Regulatory Compliance",
        "Risk Assessment",
        "Implementation Guidelines",
        "Monitoring Framework",
        "Governance Structure",
        "Documentation Standards",
        "Change Management",
        "Appendices",
        "References",
        "Glossary",
    ]

    # Simulated content generation per section
    generated = []
    for title in doc_sections:
        content = f"""
        This section titled '{title}' has been auto-generated using code analysis and uploaded configurations.
        - Code Insight: {code_text[:200]}...
        - Comments Included: {user_comments[:200]}...
        - This content follows regulatory best practices and model governance standards.
        """
        generated.append({"title": title, "content": content.strip()})

    return generated
